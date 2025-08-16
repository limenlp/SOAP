import os
import docx
from pypdf import PdfReader
import logging
import os
import autogen
import autogen.runtime_logging
import pandas as pd
import sqlite3
import json
import re

def generate_conversations(dir_name):
    file_list=os.listdir(dir_name)
    conversation_list=[]
    conversation_name=[]
    file_list=file_list[2:]
    for i in file_list:
        dir_= dir_name+i
        for j in os.listdir(dir_):
            if j!='.DS_Store' and j!='About.docx':
                conversation_path=dir_+'/'+j
                conversation_name.append(j)
                doc=docx.Document(conversation_path)
                conversation_text=[]
                for para in doc.paragraphs:
                    conversation_text.append(para.text)
                conversation='\n'.join(conversation_text)
                conversation_list.append(conversation)
    return conversation_list,conversation_name

def process_pdf(file_name):
    reader = PdfReader(file_name)
    demo_text=''
    for i in range(1,len(reader.pages)):
        page=reader.pages[i]
        text=page.extract_text()
        text=text.replace('\n','')
        demo_text+=text
    return demo_text

def rubric_processor(file):
    doc=docx.Document(file)
    rubric_content=[]
    for i in doc.paragraphs:
        rubric_content.append(i.text)
    rubric_content=rubric_content[9:-23]
    rubric="\n".join(rubric_content)
    return rubric

def conversation_db_processor(conversation_db):
    table="events"
    query=f"SELECT * from {table}"
    cursor=conversation_db.execute(query)
    rows=cursor.fetchall()
    column_names = [description[0] for description in cursor.description]
    data = [dict(zip(column_names, row)) for row in rows]
    conversation_db.close()
    soapnotes_data_df=pd.DataFrame(data)
    curr_soap_notes=''
    for j in range(0,len(soapnotes_data_df['json_state'])):
        soap_notes_dict=json.loads(soapnotes_data_df['json_state'][j])
        if '"message":' in soapnotes_data_df['json_state'][j] and '"sender":' in soapnotes_data_df['json_state'][j]:
            soap_notes_dict=json.loads(soapnotes_data_df['json_state'][j])
            if 'clinician' in soap_notes_dict['sender']:
                curr_soap_notes+='\n'+soap_notes_dict['message']
    return curr_soap_notes


def score_processor(evaluation):
    pattern1=r'(?:(?:FINAL|Final)\s*)?(?:(?:SOAP\s*NOTE|Soap\s*Note)\s*)?(?:Rating|Score|SCORE|RATING):\**\s*\**(\d+(?:\.\d+)?)'
    pattern2= r'(?:SOAP\s*Note\s*Score:\s*)?\*{0,2}(?:Rating|Score|SCORE|RATING)\*{0,2}:\s*\*{0,2}(\d+(?:\.\d+)?)(?:/\d+(?:\.\d+)?)?'
    pattern3 = r'(?:(?:FINAL|Final)\s*)?(?:Rating|Score|SCORE|RATING)\s*=\s*(\d+(?:\.\d+)?)(?:/\d+(?:\.\d+)?)?'
    pattern4 = r'(?:#+\s*)?(?:FINAL|Final)?\s*RATING\s*\n\s*(\d+(?:\.\d+)?)(?:/\d+(?:\.\d+)?)?'
    pattern5 = r'is\s+\*{0,2}(\d+(?:\.\d+)?)/\d+(?:\.\d+)?\*{0,2}'
    if(re.search(pattern1,evaluation)):
        score=float(re.findall(pattern1,evaluation)[-1])
    elif(re.search(pattern2,evaluation)):
        score=float(re.findall(pattern2,evaluation)[-1])
    elif(re.search(pattern3,evaluation)):
        score=float(re.findall(pattern3,evaluation)[-1])
    elif(re.search(pattern4,evaluation)):
        score=float(re.findall(pattern4,evaluation)[-1])
    elif(re.search(pattern5,evaluation)):
        score=float(re.findall(pattern5,evaluation)[-1])
    return score